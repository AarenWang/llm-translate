"""DOCX exporter tests."""

import tempfile
import unittest
from pathlib import Path
from docx import Document

from llm_translate.exporter_docx import DocxExporter
from llm_translate.domain import DocumentBlock, TranslationChunk, ChunkStatus


class DocxExporterTest(unittest.TestCase):
    """Test cases for DOCX exporter."""

    def setUp(self):
        """Set up test fixtures."""
        self.exporter = DocxExporter()
        self.temp_dir = tempfile.mkdtemp()

    def _create_test_docx(self, path: Path) -> None:
        """Create a test DOCX file."""
        doc = Document()
        doc.add_heading('Test Document', level=1)
        doc.add_paragraph('This is the first paragraph.')
        doc.add_paragraph('This is the second paragraph.')
        doc.add_heading('Section 1', level=2)
        doc.add_paragraph('Content for section 1.')
        doc.save(str(path))

    def _create_test_blocks_and_chunks(self):
        """Create test blocks and chunks."""
        blocks = [
            DocumentBlock(
                id="block_1",
                project_id="test_project",
                parent_id=None,
                block_order=0,
                block_type="docx_heading_1",
                level=1,
                source_text="Test Document",
            ),
            DocumentBlock(
                id="block_2",
                project_id="test_project",
                parent_id=None,
                block_order=1,
                block_type="docx_paragraph",
                level=None,
                source_text="This is the first paragraph.",
            ),
            DocumentBlock(
                id="block_3",
                project_id="test_project",
                parent_id=None,
                block_order=2,
                block_type="docx_paragraph",
                level=None,
                source_text="This is the second paragraph.",
            ),
        ]

        chunks = [
            TranslationChunk(
                id="chunk_1",
                project_id="test_project",
                chapter_id=None,
                chunk_order=0,
                block_ids=["block_1"],
                source_text="Test Document",
                target_text="测试文档",
                restored_text="测试文档",
                status=ChunkStatus.DONE,
            ),
            TranslationChunk(
                id="chunk_2",
                project_id="test_project",
                chapter_id=None,
                chunk_order=1,
                block_ids=["block_2", "block_3"],
                source_text="This is the first paragraph.\n\nThis is the second paragraph.",
                target_text="这是第一段。\n\n这是第二段。",
                restored_text="这是第一段。\n\n这是第二段。",
                status=ChunkStatus.DONE,
            ),
        ]

        return blocks, chunks

    def test_export_to_docx(self):
        """Test exporting to DOCX format."""
        # Create source DOCX
        source_path = Path(self.temp_dir) / "source.docx"
        self._create_test_docx(source_path)

        # Create test data
        blocks, chunks = self._create_test_blocks_and_chunks()

        # Create artifact directory
        artifact_dir = Path(self.temp_dir) / "artifacts"

        # Export
        result = self.exporter.export(
            artifact_dir=artifact_dir,
            source_docx_path=source_path,
            blocks=blocks,
            chunks=chunks,
            reports=[],
            draft=False
        )

        # Verify files were created
        self.assertIn("translated_docx", result)
        self.assertIn("translated_markdown", result)
        self.assertIn("bilingual", result)

        translated_docx = result["translated_docx"]
        self.assertTrue(translated_docx.exists())

        # Verify translated DOCX content
        doc = Document(str(translated_docx))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Check that translations are present
        self.assertIn("测试文档", paragraphs)

    def test_translation_map_building(self):
        """Test building translation map."""
        blocks, chunks = self._create_test_blocks_and_chunks()

        translation_map = self.exporter._build_translation_map(blocks, chunks, draft=False)

        # Verify translations are mapped correctly
        self.assertEqual(translation_map.get("block_1"), "测试文档")
        self.assertIn("第一段", translation_map.get("block_2", ""))

    def test_markdown_export(self):
        """Test Markdown export."""
        blocks, chunks = self._create_test_blocks_and_chunks()

        markdown = self.exporter._build_markdown(blocks, chunks, draft=False)

        # Verify markdown content
        self.assertIn("测试文档", markdown)
        self.assertIn("#", markdown)  # Heading markers

    def test_bilingual_export(self):
        """Test bilingual export."""
        blocks, chunks = self._create_test_blocks_and_chunks()

        bilingual = self.exporter._build_bilingual(blocks, chunks)

        # Verify bilingual content
        self.assertIn("Test Document", bilingual)
        self.assertIn("测试文档", bilingual)
        self.assertIn("Source:", bilingual)
        self.assertIn("Translation:", bilingual)

    def test_draft_mode(self):
        """Test draft mode export."""
        blocks, chunks = self._create_test_blocks_and_chunks()

        # Set both chunks to pending status
        for chunk in chunks:
            chunk.status = ChunkStatus.PENDING
            chunk.target_text = None
            chunk.restored_text = None

        translation_map = self.exporter._build_translation_map(blocks, chunks, draft=True)

        # In draft mode, pending chunks should show status
        self.assertIn("PENDING", translation_map.get("block_1", ""))
        self.assertIn("PENDING", translation_map.get("block_2", ""))

    def tearDown(self):
        """Clean up test fixtures."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
