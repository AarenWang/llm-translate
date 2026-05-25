"""DOCX protection engine and chunking tests."""

import tempfile
import unittest
from pathlib import Path
from docx import Document

from llm_translate.formats.docx import DocxFormatAdapter
from llm_translate.domain import TranslationProject


class DocxProtectionChunkingTest(unittest.TestCase):
    """Test cases for DOCX protection and chunking."""

    def setUp(self):
        """Set up test fixtures."""
        self.adapter = DocxFormatAdapter()
        self.temp_dir = tempfile.mkdtemp()

    def _create_test_docx(self, path: Path) -> None:
        """Create a test DOCX file with various elements."""
        doc = Document()

        # Add title
        doc.add_heading('Test Document', level=1)

        # Add paragraph with URL (should be protected)
        doc.add_paragraph('Visit https://example.com for more info.')

        # Add normal paragraph
        doc.add_paragraph('This is a normal paragraph for translation.')

        # Add section heading
        doc.add_heading('Section 1', level=2)

        # Add paragraph with code-like content (should be protected)
        doc.add_paragraph('Use data = {"key": "value"} for processing.')

        # Add more content
        doc.add_paragraph('Regular content follows here.')

        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Header 1'
        table.rows[0].cells[1].text = 'Header 2'
        table.rows[1].cells[0].text = 'Data 1'
        table.rows[1].cells[1].text = 'Data 2'

        doc.save(str(path))

    def test_protection_engine(self):
        """Test that protection engine protects appropriate content."""
        test_text = "Visit https://example.com and check [link](https://other.com) for info."

        # Apply protection
        result = self.adapter.protection_engine.protect("test_project", "test_chunk", test_text)

        # Verify protection worked - should detect URL and link target
        self.assertIn("__LT_", result.protected_text)
        self.assertGreater(len(result.spans), 0)  # At least one span (URL)

        # Verify restoration
        restored = self.adapter.protection_engine.restore(result.protected_text, result.spans)
        self.assertEqual(restored, test_text)

    def test_code_filtering_in_parser(self):
        """Test that code-like content is filtered by parser, not protection engine."""
        # Code-like content should be filtered by parser during extraction
        # This is verified in the parser tests, so we just need to confirm
        # the protection engine handles structural elements

        test_text = "Normal text with https://example.com URL."
        result = self.adapter.protection_engine.protect("test_project", "test_chunk", test_text)

        # URL should be protected
        self.assertIn("__LT_URL_", result.protected_text)
        self.assertEqual(len(result.spans), 1)
        self.assertEqual(result.spans[0].span_type, "URL")

    def test_chunking_strategy(self):
        """Test that chunking strategy properly groups paragraphs."""
        # Create a test DOCX file
        docx_path = Path(self.temp_dir) / "test.docx"
        self._create_test_docx(docx_path)

        # Create a mock project
        project = TranslationProject(
            id="test_project",
            name="Test Project",
            source_file_name=docx_path.name,
            source_language="en",
            target_language="zh",
            input_format="docx",
            status="CREATED"
        )

        # Parse the document
        from llm_translate.formats.base import FormatContext
        context = FormatContext(
            project_dir=Path(self.temp_dir),
            artifact_dir=Path(self.temp_dir) / "artifacts",
            source_path=docx_path,
            snapshot_dir=Path(self.temp_dir) / "snapshots"
        )

        blocks = self.adapter.parse(project, context)

        # Verify blocks were created
        self.assertGreater(len(blocks), 0)

        # Plan chunks
        chunks = self.adapter.plan_chunks("test_project", blocks)

        # Verify chunks were created
        self.assertGreater(len(chunks), 0)

        # Verify chunk properties
        for chunk in chunks:
            self.assertIsNotNone(chunk.id)
            self.assertEqual(chunk.project_id, "test_project")
            self.assertIsNotNone(chunk.source_text)

            # Check that protection was applied
            if chunk.metadata and "protected_spans" in chunk.metadata:
                self.assertIsInstance(chunk.metadata["protected_spans"], list)

    def test_format_detection(self):
        """Test format detection capabilities."""
        self.assertTrue(self.adapter.supports(Path("test.docx")))
        self.assertFalse(self.adapter.supports(Path("test.pdf")))
        self.assertFalse(self.adapter.supports(Path("test.txt")))

    def test_prompt_document_format(self):
        """Test prompt generation."""
        prompt = self.adapter.prompt_document_format()
        self.assertIn("Microsoft Word", prompt)
        self.assertIn("DOCX", prompt)
        self.assertIn("Paragraph", prompt)  # Capital P in prompt

    def tearDown(self):
        """Clean up test fixtures."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
