"""End-to-end DOCX translation pipeline test."""

import tempfile
import unittest
from pathlib import Path
from docx import Document

from llm_translate.formats.docx import DocxFormatAdapter
from llm_translate.domain import TranslationProject, ChunkStatus


class DocxPipelineTest(unittest.TestCase):
    """Test complete DOCX translation pipeline."""

    def setUp(self):
        """Set up test fixtures."""
        self.adapter = DocxFormatAdapter()
        self.temp_dir = tempfile.mkdtemp()

    def _create_test_docx(self, path: Path) -> None:
        """Create a comprehensive test DOCX file."""
        doc = Document()

        # Title
        doc.add_heading('Test Document for Translation', level=1)

        # Introduction
        doc.add_paragraph('This is a test document for the DOCX translation pipeline.')
        doc.add_paragraph('It contains various elements like paragraphs, headings, and tables.')

        # Section 1
        doc.add_heading('Introduction Section', level=2)
        doc.add_paragraph('This section demonstrates the translation capabilities.')

        # Content with URL (should be protected)
        doc.add_paragraph('Visit https://example.com for more information.')

        # Table
        doc.add_heading('Data Table', level=3)
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Column 1'
        table.rows[0].cells[1].text = 'Column 2'
        table.rows[1].cells[0].text = 'Data 1'
        table.rows[1].cells[1].text = 'Data 2'

        # Conclusion
        doc.add_heading('Conclusion', level=2)
        doc.add_paragraph('This concludes our test document.')

        doc.save(str(path))

    def test_end_to_end_pipeline(self):
        """Test the complete translation pipeline from parsing to export."""
        # Create source DOCX
        source_path = Path(self.temp_dir) / "source.docx"
        self._create_test_docx(source_path)

        # Create project
        project = TranslationProject(
            id="test_project",
            name="Test DOCX Translation",
            source_file_name=source_path.name,
            source_language="en",
            target_language="zh",
            input_format="docx",
            status="CREATED"
        )

        # Create context
        from llm_translate.formats.base import FormatContext
        project_dir = Path(self.temp_dir)
        context = FormatContext(
            project_dir=project_dir,
            artifact_dir=project_dir / "artifacts",
            source_path=source_path,
            snapshot_dir=project_dir / "snapshots"
        )

        # Step 1: Parse document
        blocks = self.adapter.parse(project, context)
        self.assertGreater(len(blocks), 0, "Should extract blocks from DOCX")
        print(f"PASS Parsed {len(blocks)} blocks from DOCX")

        # Step 2: Plan chunks
        chunks = self.adapter.plan_chunks(project.id, blocks)
        self.assertGreater(len(chunks), 0, "Should create chunks for translation")
        print(f"PASS Created {len(chunks)} chunks for translation")

        # Step 3: Simulate translation
        for chunk in chunks:
            chunk.target_text = f"[TRANSLATED: {chunk.source_text[:50]}...]"
            chunk.restored_text = chunk.target_text
            chunk.status = ChunkStatus.DONE
        print("PASS Simulated translation for all chunks")

        # Step 4: Export translated content
        result_paths, reports, has_changes = self.adapter.export(
            project, context, blocks, chunks, [], draft=False
        )
        self.assertTrue(has_changes, "Should detect changes")
        self.assertIn("translated_docx", result_paths, "Should export DOCX")
        print(f"PASS Exported {len(result_paths)} files")

        # Verify outputs exist
        for output_type, path in result_paths.items():
            self.assertTrue(path.exists(), f"{output_type} file should exist")
            print(f"  PASS {output_type}: {path.name}")

        # Step 5: Validate exported document
        translated_docx = result_paths["translated_docx"]
        doc = Document(str(translated_docx))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        self.assertGreater(len(paragraphs), 0, "Translated DOCX should have content")
        print(f"PASS Translated DOCX contains {len(paragraphs)} paragraphs")

        # Check that translation markers are present
        all_text = " ".join(paragraphs)
        self.assertIn("TRANSLATED", all_text, "Should contain translation markers")
        print("PASS Translation content verified in exported DOCX")

    def tearDown(self):
        """Clean up test fixtures."""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
