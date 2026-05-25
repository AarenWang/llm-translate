"""DOCX parser tests."""

import tempfile
import zipfile
from pathlib import Path
import unittest

from docx import Document
from docx.shared import Inches, Pt

from llm_translate.parser.docx import DocxParser


class DocxParserTest(unittest.TestCase):
    """Test cases for DOCX parser."""

    def setUp(self):
        """Set up test fixtures."""
        self.parser = DocxParser()

    def _create_simple_docx(self, path: Path) -> None:
        """Create a simple DOCX file for testing."""
        doc = Document()

        # Add heading
        doc.add_heading('Test Document', level=1)

        # Add normal paragraphs
        doc.add_paragraph('This is the first paragraph.')
        doc.add_paragraph('This is the second paragraph.')

        # Add heading
        doc.add_heading('Section 1', level=2)

        # Add more content
        doc.add_paragraph('Content for section 1.')

        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Column 1'
        table.rows[0].cells[1].text = 'Column 2'
        table.rows[1].cells[0].text = 'Data 1'
        table.rows[1].cells[1].text = 'Data 2'

        # Add list
        doc.add_paragraph('List item 1', style='List Bullet')
        doc.add_paragraph('List item 2', style='List Bullet')

        doc.save(str(path))

    def _create_complex_docx(self, path: Path) -> None:
        """Create a complex DOCX file for testing."""
        doc = Document()

        # Title
        doc.add_heading('Complex Document Test', level=1)

        # Abstract section
        doc.add_paragraph('This is an abstract paragraph.')

        # Main content
        doc.add_heading('Introduction', level=2)
        doc.add_paragraph('This is the introduction paragraph with some content.')
        doc.add_paragraph('Another paragraph for testing.')

        # Subsection
        doc.add_heading('Background', level=3)
        doc.add_paragraph('Background information goes here.')

        # Another section
        doc.add_heading('Methods', level=2)
        doc.add_paragraph('Method description goes here.')

        # Save
        doc.save(str(path))

    def test_parse_simple_docx(self) -> None:
        """Test parsing a simple DOCX file."""
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "simple.docx"
            self._create_simple_docx(docx_path)

            # Parse the document
            result = self.parser.parse("test_project", docx_path)

            # Verify basic results
            self.assertGreater(len(result.blocks), 0)
            self.assertGreater(len(result.paragraphs), 0)
            self.assertEqual(len(result.tables), 1)
            self.assertGreater(len(result.heading_structure), 0)

    def test_extract_headings(self) -> None:
        """Test heading extraction."""
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "headings.docx"
            self._create_complex_docx(docx_path)

            result = self.parser.parse("test_project", docx_path)

            # Verify headings
            headings = result.heading_structure
            self.assertGreater(len(headings), 0)

            # Check heading levels
            levels = [h['level'] for h in headings]
            self.assertIn(1, levels)  # Should have H1
            self.assertIn(2, levels)  # Should have H2
            self.assertIn(3, levels)  # Should have H3

    def test_paragraph_extraction(self) -> None:
        """Test paragraph extraction."""
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "paragraphs.docx"
            self._create_simple_docx(docx_path)

            result = self.parser.parse("test_project", docx_path)

            # Verify paragraphs
            paragraphs = result.paragraphs
            self.assertGreater(len(paragraphs), 0)

            # Check translatable paragraphs
            translatable = [p for p in paragraphs if p.is_translatable]
            self.assertGreater(len(translatable), 0)

    def test_table_extraction(self) -> None:
        """Test table extraction."""
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "tables.docx"
            self._create_simple_docx(docx_path)

            result = self.parser.parse("test_project", docx_path)

            # Verify tables
            tables = result.tables
            self.assertEqual(len(tables), 1)

            # Check table structure
            table = tables[0]
            self.assertEqual(table.rows, 2)
            self.assertEqual(table.cols, 2)
            self.assertEqual(len(table.cells), 4)

    def test_metadata_extraction(self) -> None:
        """Test metadata extraction."""
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "metadata.docx"
            self._create_simple_docx(docx_path)

            result = self.parser.parse("test_project", docx_path)

            # Verify metadata
            metadata = result.metadata
            self.assertIn("total_paragraphs", metadata)
            self.assertIn("total_tables", metadata)
            self.assertIn("total_headings", metadata)
            self.assertIn("styles_used", metadata)

    def test_heading_level_detection(self) -> None:
        """Test heading level detection."""
        # Test the _get_heading_level method
        test_cases = [
            ("Heading 1", 1),
            ("Heading 2", 2),
            ("Heading 6", 6),
            ("Normal", None),
            ("Title", None),
        ]

        for style_name, expected_level in test_cases:
            result = self.parser._get_heading_level(style_name)
            self.assertEqual(result, expected_level, f"Failed for style: {style_name}")

    def test_code_detection(self) -> None:
        """Test code-like content detection."""
        # Test the _looks_like_code method
        code_like_cases = [
            ("def function(): { return value; }", False),  # Not enough code indicators
            ("int main(int argc, char** argv) { return 0; }", False),  # Not enough parentheses
            ("for i in range(10): print(i);", True),  # Code-like with colon and parentheses
            ("data = {'key': 'value'}; process(data);", True),  # Code with braces and semicolons
            ("This is normal text.", False),  # Normal text
            ("Some text with numbers: 1, 2, 3.", False),  # Normal text
        ]

        for text, expected_result in code_like_cases:
            result = self.parser._looks_like_code(text)
            self.assertEqual(result, expected_result, f"Failed for text: {text}")

    def test_parse_empty_docx(self) -> None:
        """Test parsing an empty DOCX file."""
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "empty.docx"
            doc = Document()  # Create empty document
            doc.save(str(docx_path))

            result = self.parser.parse("test_project", docx_path)

            # Empty document should have some blocks (metadata)
            self.assertEqual(len(result.blocks), 0)
            self.assertEqual(len(result.paragraphs), 0)

    def test_parse_nonexistent_file(self) -> None:
        """Test parsing a non-existent file."""
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "nonexistent.docx"

            with self.assertRaises(FileNotFoundError):
                self.parser.parse("test_project", docx_path)


if __name__ == "__main__":
    unittest.main()