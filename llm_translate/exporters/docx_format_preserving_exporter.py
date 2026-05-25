"""Format-preserving PDF exporter using PDF→DOCX→Translate→DOCX→PDF conversion chain."""

from __future__ import annotations

import tempfile
import shutil
from pathlib import Path
from typing import Any

from ..domain import TranslationProject, TranslationChunk, ValidationReport
from ..formats.base import FormatContext


class DocxFormatPreservingExporter:
    """Export translated PDF while preserving format via DOCX conversion."""

    def export(
        self,
        project: TranslationProject,
        context: FormatContext,
        blocks: list[Any],
        chunks: list[TranslationChunk],
        reports: list[ValidationReport],
        draft: bool,
    ) -> Path:
        """Export as format-preserved PDF using DOCX conversion chain."""
        output_path = context.artifact_dir / f"{project.name}_format_preserved.pdf"

        # Ensure artifacts directory exists
        context.artifact_dir.mkdir(parents=True, exist_ok=True)

        try:
            # Step 1: Convert PDF to DOCX
            print("[FORMAT_PRESERVE] Step 1: Converting PDF to DOCX...")
            docx_path = self._convert_pdf_to_docx(context.source_path, context.artifact_dir)

            # Step 2: Create translated DOCX
            print("[FORMAT_PRESERVE] Step 2: Creating translated DOCX...")
            translated_docx_path = self._create_translated_docx(
                docx_path, chunks, context.artifact_dir, project.name
            )

            # Step 3: Convert translated DOCX back to PDF
            print("[FORMAT_PRESERVE] Step 3: Converting translated DOCX to PDF...")
            self._convert_docx_to_pdf(translated_docx_path, output_path)

            print(f"[FORMAT_PRESERVE] Format-preserved PDF created: {output_path}")

        except Exception as e:
            print(f"[FORMAT_PRESERVE_ERROR] {e}")
            print("[FORMAT_PRESERVE_WARN] Format preservation failed, falling back to simple PDF")
            # Fall back to simple PDF generation
            from .pdf_exporters import SimplePdfExporter

            fallback_exporter = SimplePdfExporter()
            return fallback_exporter.export(
                project, context, blocks, chunks, reports, draft
            )

        return output_path

    def _convert_pdf_to_docx(self, pdf_path: Path, output_dir: Path) -> Path:
        """Convert PDF to DOCX while preserving format."""
        try:
            from pdf2docx import Converter

            docx_path = output_dir / "temp_original.docx"
            cv = Converter(str(pdf_path))

            # Convert PDF to DOCX
            cv.convert(str(docx_path))
            cv.close()

            return docx_path

        except ImportError:
            raise ImportError(
                "pdf2docx is required for format preservation. "
                "Install with: pip install pdf2docx"
            )
        except Exception as e:
            raise Exception(f"PDF to DOCX conversion failed: {e}")

    def _create_translated_docx(
        self, docx_path: Path, chunks: list[TranslationChunk], output_dir: Path, project_name: str
    ) -> Path:
        """Create translated DOCX by replacing text."""
        try:
            from docx import Document

            # Load original DOCX
            doc = Document(docx_path)

            # Build translation map
            translation_map = self._build_translation_map(chunks)

            # Replace text in paragraphs
            for paragraph in doc.paragraphs:
                original_text = paragraph.text.strip()
                if original_text and original_text in translation_map:
                    # Replace text while preserving formatting
                    translated_text = translation_map[original_text]
                    # Clear paragraph and add translated text
                    paragraph.clear()
                    paragraph.add_run(translated_text)

            # Replace text in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text.strip()
                            if original_text and original_text in translation_map:
                                translated_text = translation_map[original_text]
                                paragraph.clear()
                                paragraph.add_run(translated_text)

            # Save translated DOCX
            translated_docx_path = output_dir / f"{project_name}_translated.docx"
            doc.save(translated_docx_path)

            return translated_docx_path

        except Exception as e:
            raise Exception(f"DOCX translation failed: {e}")

    def _convert_docx_to_pdf(self, docx_path: Path, pdf_path: Path) -> None:
        """Convert DOCX to PDF."""
        try:
            # Try using docx2pdf
            try:
                from docx2pdf import convert

                convert(docx_path, pdf_path)
                return

            except ImportError:
                pass

            # Fallback: use LibreOffice if available
            import subprocess
            import platform

            system = platform.system()
            try:
                if system == "Windows":
                    # Try LibreOffice on Windows
                    subprocess.run([
                        "soffice",
                        "--headless",
                        "--convert-to", "pdf",
                        "--outdir", str(pdf_path.parent),
                        str(docx_path)
                    ], check=True, capture_output=True)
                    return
                elif system == "Darwin":  # macOS
                    subprocess.run([
                        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                        "--headless",
                        "--convert-to", "pdf",
                        "--outdir", str(pdf_path.parent),
                        str(docx_path)
                    ], check=True, capture_output=True)
                    return
            except (subprocess.CalledProcessError, FileNotFoundError):
                pass

            # Last resort: manual PDF generation with reportlab
            print("[DOCX2PDF_WARN] Automatic conversion failed, using reportlab fallback")
            self._docx_to_pdf_with_reportlab(docx_path, pdf_path)

        except Exception as e:
            raise Exception(f"DOCX to PDF conversion failed: {e}")

    def _docx_to_pdf_with_reportlab(self, docx_path: Path, pdf_path: Path) -> None:
        """Fallback DOCX to PDF conversion using reportlab."""
        try:
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

            # Load DOCX
            doc = Document(docx_path)

            # Create PDF
            pdf_doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Add content from DOCX
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    style = styles["BodyText"]
                    story.append(Paragraph(paragraph.text, style))
                    story.append(Spacer(1, 6))

            pdf_doc.build(story)

        except Exception as e:
            raise Exception(f"Reportlab fallback conversion failed: {e}")

    def _build_translation_map(self, chunks: list[TranslationChunk]) -> dict[str, str]:
        """Build mapping from source text to translated text."""
        translation_map = {}

        for chunk in chunks:
            if chunk.source_text and chunk.restored_text:
                # Split into paragraphs for more precise matching
                source_paragraphs = chunk.source_text.split('\n\n')
                target_paragraphs = chunk.restored_text.split('\n\n')

                for src, tgt in zip(source_paragraphs, target_paragraphs):
                    src, tgt = src.strip(), tgt.strip()
                    if src and tgt and len(src) > 10:  # Only match substantial text
                        translation_map[src] = tgt

                # Also add full chunk mapping
                translation_map[chunk.source_text.strip()] = chunk.restored_text.strip()

        return translation_map
