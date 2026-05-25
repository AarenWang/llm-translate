"""DOCX exporter for translated documents."""

from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any

from docx import Document

from .domain import DocumentBlock, TranslationChunk, ValidationReport
from .protection import ProtectionEngine
from .exporter import MarkdownExporter


class DocxExporter:
    """Export translated content back to DOCX format."""

    def __init__(self):
        self.protection_engine = ProtectionEngine()
        self.markdown_exporter = MarkdownExporter()

    def export(
        self,
        artifact_dir: Path,
        source_docx_path: Path,
        blocks: list[DocumentBlock],
        chunks: list[TranslationChunk],
        reports: list[ValidationReport],
        draft: bool = False,
    ) -> dict[str, Path]:
        """Export translated content to DOCX format.

        Args:
            artifact_dir: Directory to save exported files
            source_docx_path: Original DOCX file path
            blocks: Original document blocks
            chunks: Translated chunks
            reports: Validation reports
            draft: Whether this is a draft export

        Returns:
            Dictionary mapping export types to file paths
        """
        artifact_dir.mkdir(parents=True, exist_ok=True)
        suffix = ".draft" if draft else ""

        # Create output paths
        translated_docx_path = artifact_dir / f"translated{suffix}.docx"
        markdown_path = artifact_dir / f"translated{suffix}.md"
        bilingual_path = artifact_dir / f"bilingual{suffix}.md"
        log_path = artifact_dir / "translation-log.json"
        report_json_path = artifact_dir / "validation-report.json"
        report_md_path = artifact_dir / "validation-report.md"

        # Export to DOCX format
        self._export_to_docx(
            translated_docx_path,
            source_docx_path,
            blocks,
            chunks,
            draft
        )

        # Export to Markdown for review
        markdown_path.write_text(self._build_markdown(blocks, chunks, draft), encoding="utf-8")
        bilingual_path.write_text(self._build_bilingual(blocks, chunks), encoding="utf-8")

        # Write logs and reports
        log_path.write_text(
            json.dumps([self._chunk_log(chunk) for chunk in chunks], ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        report_json_path.write_text(
            json.dumps([report.__dict__ for report in reports], ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        report_md_path.write_text(self._report_markdown(reports), encoding="utf-8")

        return {
            "translated_docx": translated_docx_path,
            "translated_markdown": markdown_path,
            "bilingual": bilingual_path,
            "translation_log": log_path,
            "validation_report_json": report_json_path,
            "validation_report_md": report_md_path,
        }

    def _export_to_docx(
        self,
        output_path: Path,
        source_path: Path,
        blocks: list[DocumentBlock],
        chunks: list[TranslationChunk],
        draft: bool,
    ) -> None:
        """Export translated content to DOCX format.

        Args:
            output_path: Path to save translated DOCX
            source_path: Original DOCX file path
            blocks: Original document blocks
            chunks: Translated chunks
            draft: Whether this is a draft export
        """
        # Copy original document to preserve formatting
        shutil.copy(source_path, output_path)
        doc = Document(str(output_path))

        # Create a mapping of block IDs to translated text
        translation_map = self._build_translation_map(blocks, chunks, draft)

        # Update paragraphs with translated content
        paragraph_index = 0
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph_index >= len(blocks):
                break

            block = blocks[paragraph_index]
            if block.id in translation_map:
                translated_text = translation_map[block.id]
                if translated_text:
                    # Clear existing text and add translated text
                    paragraph.clear()
                    paragraph.add_run(translated_text)

            paragraph_index += 1

        # Update table cells with translated content
        table_index = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if table_index < len(blocks):
                        block = blocks[table_index]
                        if block.id in translation_map:
                            translated_text = translation_map[block.id]
                            if translated_text:
                                cell.text = translated_text
                        table_index += 1

        # Save the modified document
        doc.save(str(output_path))

    def _build_translation_map(
        self,
        blocks: list[DocumentBlock],
        chunks: list[TranslationChunk],
        draft: bool,
    ) -> dict[str, str]:
        """Build a mapping from block IDs to translated text.

        Args:
            blocks: Original document blocks
            chunks: Translated chunks
            draft: Whether this is a draft export

        Returns:
            Dictionary mapping block IDs to translated text
        """
        translation_map = {}

        # Create block-to-chunk mapping
        block_to_chunks: dict[str, list[TranslationChunk]] = {}
        for chunk in chunks:
            for block_id in chunk.block_ids:
                if block_id not in block_to_chunks:
                    block_to_chunks[block_id] = []
                block_to_chunks[block_id].append(chunk)

        # Process each block
        for block in blocks:
            if block.id not in block_to_chunks:
                continue

            related_chunks = block_to_chunks[block.id]
            if not related_chunks:
                continue

            # Get the primary chunk for this block
            primary_chunk = related_chunks[0]

            # Restore protected content if available
            if primary_chunk.restored_text:
                translated_text = primary_chunk.restored_text
            elif primary_chunk.target_text:
                # Restore protected content
                protected_spans = primary_chunk.metadata.get("protected_spans", [])
                if protected_spans:
                    from .domain import ProtectedSpan
                    spans = [ProtectedSpan(**span) for span in protected_spans]
                    translated_text = self.protection_engine.restore(
                        primary_chunk.target_text, spans
                    )
                else:
                    translated_text = primary_chunk.target_text
            else:
                # No translation available
                if draft:
                    translated_text = f"[{primary_chunk.status.value}]"
                else:
                    translated_text = block.source_text

            translation_map[block.id] = translated_text

        return translation_map

    def _build_markdown(
        self,
        blocks: list[DocumentBlock],
        chunks: list[TranslationChunk],
        draft: bool,
    ) -> str:
        """Build Markdown representation of translated content.

        Args:
            blocks: Original document blocks
            chunks: Translated chunks
            draft: Whether this is a draft export

        Returns:
            Markdown formatted text
        """
        parts = []

        for block in blocks:
            # Get translated text for this block
            translated_text = self._get_block_translation(block, chunks, draft)

            if not translated_text:
                continue

            # Add heading markers
            if block.block_type.startswith("docx_heading"):
                level = block.metadata.get("heading_level", 1)
                heading_prefix = "#" * level
                parts.append(f"{heading_prefix} {translated_text}")
            elif block.block_type == "docx_table_cell":
                parts.append(f"| {translated_text} |")
            else:
                parts.append(translated_text)

        return "\n\n".join(parts) + "\n"

    def _build_bilingual(
        self,
        blocks: list[DocumentBlock],
        chunks: list[TranslationChunk],
    ) -> str:
        """Build bilingual comparison document.

        Args:
            blocks: Original document blocks
            chunks: Translated chunks

        Returns:
            Bilingual comparison text
        """
        parts = []

        for block in blocks:
            # Get translated text for this block
            translated_text = self._get_block_translation(block, chunks, False)

            if not translated_text:
                translated_text = "[NOT TRANSLATED]"

            parts.append(f"## Source: {block.id}\n\n{block.source_text}\n\n")
            parts.append(f"## Translation: {block.id}\n\n{translated_text}\n\n")

        return "\n---\n\n".join(parts)

    def _get_block_translation(
        self,
        block: DocumentBlock,
        chunks: list[TranslationChunk],
        draft: bool,
    ) -> str | None:
        """Get translated text for a specific block.

        Args:
            block: Document block
            chunks: Translated chunks
            draft: Whether this is a draft export

        Returns:
            Translated text or None
        """
        # Find chunks that contain this block
        for chunk in chunks:
            if block.id in chunk.block_ids:
                if chunk.restored_text:
                    return chunk.restored_text
                elif chunk.target_text:
                    # Restore protected content
                    protected_spans = chunk.metadata.get("protected_spans", [])
                    if protected_spans:
                        from .domain import ProtectedSpan
                        spans = [ProtectedSpan(**span) for span in protected_spans]
                        return self.protection_engine.restore(chunk.target_text, spans)
                    else:
                        return chunk.target_text
                else:
                    if draft:
                        return f"[{chunk.status.value}]"

        return None

    def _chunk_log(self, chunk: TranslationChunk) -> dict:
        """Build log entry for a chunk.

        Args:
            chunk: Translation chunk

        Returns:
            Dictionary with chunk information
        """
        return {
            "chunk_id": chunk.id,
            "status": chunk.status.value,
            "retry_count": chunk.retry_count,
            "model_name": chunk.model_name,
            "prompt_version": chunk.prompt_version,
            "glossary_version": chunk.glossary_version,
            "style_guide_version": chunk.style_guide_version,
            "protection_policy_version": chunk.protection_policy_version,
            "error_message": chunk.error_message,
        }

    def _report_markdown(self, reports: list[ValidationReport]) -> str:
        """Build Markdown validation report.

        Args:
            reports: Validation reports

        Returns:
            Markdown formatted report
        """
        if not reports:
            return "# Validation Report\n\nNo validation reports.\n"

        rows = [
            "# Validation Report",
            "",
            "| Chunk | Check | Status | Issues |",
            "|---|---|---|---|"
        ]

        for report in reports:
            issues = "; ".join(issue["type"] for issue in report.issues) or "-"
            rows.append(f"| {report.chunk_id or '-'} | {report.check_type} | {report.status} | {issues} |")

        return "\n".join(rows) + "\n"
