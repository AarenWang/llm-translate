"""DOCX parser implementation."""

from __future__ import annotations

from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any

from docx import Document

from ..domain import DocumentBlock


@dataclass(frozen=True)
class DocxParagraph:
    """Represents a paragraph in a DOCX document."""
    index: int
    text: str
    style_name: str
    heading_level: int | None
    is_translatable: bool
    table_reference: str | None
    section_info: dict[str, Any]


@dataclass(frozen=True)
class DocxTable:
    """Represents a table in a DOCX document."""
    index: int
    rows: int
    cols: int
    cells: list[str]
    is_translatable: bool


@dataclass(frozen=True)
class DocxParseResult:
    """Result of parsing a DOCX document."""
    blocks: list[DocumentBlock]
    paragraphs: list[DocxParagraph]
    tables: list[DocxTable]
    heading_structure: list[dict[str, Any]]
    metadata: dict[str, Any]

    def to_snapshot(self) -> dict[str, Any]:
        """Convert parse result to snapshot for saving."""
        return {
            "paragraphs": [asdict(p) for p in self.paragraphs],
            "tables": [asdict(t) for t in self.tables],
            "heading_structure": self.heading_structure,
            "metadata": self.metadata
        }


class DocxParser:
    """Parse DOCX documents and extract translatable content."""

    # Elements that should not be translated
    SKIP_STYLES = {
        "Code",
        "Code Glyph",
        "Code Character",
        "Literal",
        "No Spacing",
        "No Proofing",
        "TOC Heading",
    }

    # Elements that indicate sections
    HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6", "Title"}

    def parse(self, project_id: str, docx_path: Path) -> DocxParseResult:
        """
        Parse a DOCX document and extract translatable content.

        Args:
            project_id: Translation project identifier
            docx_path: Path to the DOCX file

        Returns:
            DocxParseResult containing blocks, paragraphs, tables, and structure
        """
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")

        doc = Document(str(docx_path))

        # Extract document structure
        paragraphs = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)
        heading_structure = self._build_heading_structure(doc)

        # Create document blocks for translatable content
        blocks = self._create_document_blocks(project_id, paragraphs, tables)

        # Build metadata
        metadata = {
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "total_headings": len(heading_structure),
            "styles_used": self._extract_used_styles(doc),
            "has_hyperlinks": self._has_hyperlinks(doc),
            "has_images": self._has_images(doc),
            "file_name": docx_path.name,
            "file_size": docx_path.stat().st_size,
        }

        return DocxParseResult(
            blocks=blocks,
            paragraphs=paragraphs,
            tables=tables,
            heading_structure=heading_structure,
            metadata=metadata
        )

    def _extract_paragraphs(self, doc: Document) -> list[DocxParagraph]:
        """Extract paragraphs from the document."""
        paragraphs = []

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            style_name = paragraph.style.name

            # Determine heading level
            heading_level = self._get_heading_level(style_name)

            # Determine if translatable
            is_translatable = self._is_translatable_paragraph(text, style_name)

            # Check table reference
            table_reference = self._get_table_reference(paragraph)

            # Get section info
            section_info = self._get_section_info(paragraph, heading_level)

            docx_paragraph = DocxParagraph(
                index=i,
                text=text,
                style_name=style_name,
                heading_level=heading_level,
                is_translatable=is_translatable,
                table_reference=table_reference,
                section_info=section_info
            )

            paragraphs.append(docx_paragraph)

        return paragraphs

    def _extract_tables(self, doc: Document) -> list[DocxTable]:
        """Extract tables from the document."""
        tables = []

        for i, table in enumerate(doc.tables):
            cells = []

            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cells.append(cell_text)

            docx_table = DocxTable(
                index=i,
                rows=len(table.rows),
                cols=len(table.columns),
                cells=cells,
                is_translatable=True  # Table cells are translatable
            )

            tables.append(docx_table)

        return tables

    def _build_heading_structure(self, doc: Document) -> list[dict[str, Any]]:
        """Build heading hierarchy structure."""
        structure = []

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.style.name in self.HEADING_STYLES:
                heading_level = self._get_heading_level(paragraph.style.name)
                if heading_level:
                    structure.append({
                        "index": i,
                        "level": heading_level,
                        "text": paragraph.text,
                        "style": paragraph.style.name
                    })

        return structure

    def _create_document_blocks(
        self,
        project_id: str,
        paragraphs: list[DocxParagraph],
        tables: list[DocxTable]
    ) -> list[DocumentBlock]:
        """Create document blocks from paragraphs and tables."""
        blocks = []

        # Create blocks for translatable paragraphs
        for paragraph in paragraphs:
            if paragraph.is_translatable and paragraph.text:
                block_order = len(blocks)
                block = DocumentBlock(
                    id=f"{project_id}_b_{block_order + 1:06d}",
                    project_id=project_id,
                    parent_id=None,
                    block_order=block_order,
                    block_type=self._get_block_type(paragraph),
                    level=paragraph.heading_level,
                    source_text=paragraph.text,
                    metadata={
                        "format": "docx",
                        "paragraph_index": paragraph.index,
                        "style_name": paragraph.style_name,
                        "heading_level": paragraph.heading_level,
                        "table_reference": paragraph.table_reference,
                        "section_info": paragraph.section_info,
                        "translatable": True,
                    }
                )
                blocks.append(block)

        # Create blocks for table cells
        for table in tables:
            if table.is_translatable:
                for cell_index, cell_text in enumerate(table.cells):
                    if cell_text.strip():
                        block_order = len(blocks)
                        block = DocumentBlock(
                            id=f"{project_id}_b_{block_order + 1:06d}",
                            project_id=project_id,
                            parent_id=None,
                            block_order=block_order,
                            block_type="docx_table_cell",
                            level=None,
                            source_text=cell_text,
                            metadata={
                                "format": "docx",
                                "table_index": table.index,
                                "cell_index": cell_index,
                                "translatable": True,
                            }
                        )
                        blocks.append(block)

        return blocks

    def _get_heading_level(self, style_name: str) -> int | None:
        """Extract heading level from style name."""
        if style_name.startswith("Heading ") and len(style_name) > 8:
            try:
                return int(style_name[8:])
            except ValueError:
                return None
        return None

    def _is_translatable_paragraph(self, text: str, style_name: str) -> bool:
        """Determine if a paragraph should be translated."""
        if not text.strip():
            return False

        if style_name in self.SKIP_STYLES:
            return False

        # Skip if it looks like code or special content
        if self._looks_like_code(text):
            return False

        return True

    def _looks_like_code(self, text: str) -> bool:
        """Check if text looks like code."""
        # Simple heuristic: if it contains many special characters or is very uniform
        if len(text) < 10:
            return False

        # Check for patterns that look like code
        code_indicators = [
            # Dictionary/object literals with assignment
            lambda t: '{' in t and '}' in t and t.count('=') >= 1,
            # Array/object access with assignment
            lambda t: t.count('[') >= 1 and t.count(']') >= 1 and t.count('=') >= 1,
            # Multiple statements with semicolons and assignment
            lambda t: t.count(';') >= 2 and t.count('=') >= 1 and ('(' in t or ')' in t),
            # Multiple assignments or object properties
            lambda t: t.count('=') >= 2 and ('("' in t or '")' in t or ",'" in t or "')" in t),
            # Python-style code (keywords + colons + parens, without braces)
            lambda t: ':' in t and '(' in t and ')' in t and t.count('{') == 0 and t.count('}') == 0,
        ]

        return any(indicator(text) for indicator in code_indicators)

    def _get_table_reference(self, paragraph) -> str | None:
        """Check if paragraph is inside a table."""
        # This is a simplified check - a more robust implementation
        # would need to traverse the document tree
        return None

    def _get_section_info(self, paragraph, heading_level: int | None) -> dict[str, Any]:
        """Extract section information for a paragraph."""
        return {
            "heading_level": heading_level,
            "is_heading": heading_level is not None,
        }

    def _get_block_type(self, paragraph: DocxParagraph) -> str:
        """Determine the block type based on paragraph characteristics."""
        if paragraph.heading_level:
            return f"docx_heading_{paragraph.heading_level}"
        elif paragraph.table_reference:
            return "docx_table_paragraph"
        else:
            return "docx_paragraph"

    def _extract_used_styles(self, doc: Document) -> list[str]:
        """Extract list of used styles in the document."""
        styles = set()
        for paragraph in doc.paragraphs:
            styles.add(paragraph.style.name)
        return sorted(list(styles))

    def _has_hyperlinks(self, doc: Document) -> bool:
        """Check if document contains hyperlinks."""
        # Simple check for hyperlinks
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if 'http' in text or 'www.' in text:
                return True
        return False

    def _has_images(self, doc: Document) -> bool:
        """Check if document contains images."""
        # Check for image relationships
        return len(doc.part.rels) > 0